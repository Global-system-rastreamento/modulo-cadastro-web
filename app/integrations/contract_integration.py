from app.src.docx_funcs import set_table_borders

import docx.document
import requests
import streamlit as st
from datetime import date
import docx


def formatar_valor_financeiro_contrato(valor):
    if valor is None or valor == '---': return "0,00"
    try:
        return f"{float(valor):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except ValueError:
        return "0,00"

def formatar_telefone_contrato(numero_tel):
    if not numero_tel: return ""
    num_limpo = "".join(filter(str.isdigit, str(numero_tel)))
    if len(num_limpo) == 11: # (XX) XXXXX-XXXX
        return f"({num_limpo[0:2]}) {num_limpo[2:7]}-{num_limpo[7:]}"
    elif len(num_limpo) == 10: # (XX) XXXX-XXXX
        return f"({num_limpo[0:2]}) {num_limpo[2:6]}-{num_limpo[6:]}"
    return str(numero_tel)

def formatar_cep_contrato(cep):
    if not cep: return ""
    cep_limpo = "".join(filter(str.isdigit, str(cep)))
    if len(cep_limpo) == 8:
        return f"{cep_limpo[0:5]}-{cep_limpo[5:]}"
    return str(cep)

def formatar_cpf_contrato(cpf):
    if not cpf: return ""
    cpf_limpo = "".join(filter(str.isdigit, str(cpf)))
    if len(cpf_limpo) == 11:
        return f"{cpf_limpo[:3]}.{cpf_limpo[3:6]}.{cpf_limpo[6:9]}-{cpf_limpo[9:]}"
    return str(cpf)

def formatar_cnpj_contrato(cnpj):
    if not cnpj: return ""
    cnpj_limpo = "".join(filter(str.isdigit, str(cnpj)))
    if len(cnpj_limpo) == 14:
        return f"{cnpj_limpo[:2]}.{cnpj_limpo[2:5]}.{cnpj_limpo[5:8]}/{cnpj_limpo[8:12]}-{cnpj_limpo[12:]}"
    return str(cnpj)

# --- Funções de Geração de Contrato e Manual ---
def baixar_template_docx(url):
    """Baixa um template DOCX de uma URL."""
    try:
        response = requests.get(url, timeout=20)
        response.raise_for_status()
        return response.content
    except requests.exceptions.RequestException as e:
        st.error(f"Erro ao baixar template de {url}: {e}")
        return None

def preparar_dados_para_template_contrato():
    """Coleta e formata os dados do st.session_state para o template do contrato."""
    dados = {}
    meses_extenso = ['janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho', 'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro']
    hoje = date.today()
    data_instalacao_obj = st.session_state.contract_data_instalacao_input

    # Dados do Cliente/Contratante (do formulário principal)
    nome_contratante = st.session_state.form_nome.split("- ")[-1]
    cpf_cnpj_contratante = st.session_state.form_cpf_cnpj
    
    # Tenta extrair Rua, Número, Bairro, Cidade, Estado, CEP do endereço completo
    endereco_completo = st.session_state.form_endereco
    # Heurística simples para extrair partes do endereço. Pode precisar de melhorias.
    partes_endereco = [p.strip() for p in endereco_completo.split(',')]
    
    dados['CORPO'] = {
        'NOME_CLIENTE_COM_COD': st.session_state.form_nome,
        '_DATA_INSTALACAO_EXTENSO_': f"{data_instalacao_obj.day} de {meses_extenso[data_instalacao_obj.month - 1]} de {data_instalacao_obj.year}",
        '_NOME_RESPONSAVEL_': st.session_state.form_responsavel if st.session_state.form_responsavel else nome_contratante,
        '_LOCAL_INSTALACAO_': st.session_state.contract_local_instalacao_input if not st.session_state.contract_local_instalacao_input == "Outro" else st.session_state.contract_local_instalacao_outro_input,
    }
    dados['TABELAS'] = {
        '_NOME_CONTRATANTE_': nome_contratante,
        '_CPF_CNPJ_CONTRATANTE_': cpf_cnpj_contratante,
        '_NOME_RESPONSAVEL_': st.session_state.form_responsavel if st.session_state.form_responsavel else nome_contratante,
        '_RUA_': st.session_state.dados_cobranca_endereco if not st.session_state.dados_cobranca_complemento else f"{st.session_state.dados_cobranca_endereco}, {st.session_state.dados_cobranca_complemento}",
        '_BAIRRO_': st.session_state.dados_cobranca_bairro,
        '_N_': st.session_state.dados_cobranca_numero,
        '_CIDADE_': st.session_state.dados_cobranca_cidade,
        '_ESTADO_': st.session_state.dados_cobranca_estado,
        '_CEP_': formatar_cep_contrato(st.session_state.dados_cobranca_cep),
        '_FONE_1_': formatar_telefone_contrato(st.session_state.form_tel_celular or st.session_state.form_tel_comercial),
        '_FONE_2_': formatar_telefone_contrato(st.session_state.form_tel_residencial),
        '_EMAIL_': st.session_state.form_email,
        '_ADESAO_': formatar_valor_financeiro_contrato(st.session_state.contract_valor_adesao_input),
        '_DESINSTALACAO_': formatar_valor_financeiro_contrato(st.session_state.contract_valor_desinstalacao_input),
        '_REINSTALACAO_': formatar_valor_financeiro_contrato(st.session_state.contract_valor_reinstalacao_input),
        '_OPERADORA_': st.session_state.contract_operadora_input,
        '_DATA_INSTALACAO_': data_instalacao_obj.strftime('%d/%m/%Y'),
        '_VENCIMENTO_': str(st.session_state.form_dia_vencimento),
        '_MENSALIDADE_PLANO2_': formatar_valor_financeiro_contrato(st.session_state.contract_placas_list[0]['mensalidade'] if st.session_state.contract_placas_list and st.session_state.contract_tipo_contrato_select == "PLANO2" else 0), # Ajustar para pegar a mensalidade correta do plano2
        '_VALOR_TOTAL_DA_COBERTURA': formatar_valor_financeiro_contrato(st.session_state.contract_valor_cobertura_input if st.session_state.contract_cliente_seguradora_checkbox else 0),
        '_FORMA_COBRANÇA_': st.session_state.contract_forma_cobranca_input,
        '_ATENDENTE_': st.session_state.contract_atendente_input,
        '_BOOL_SEGURADORA_': "Sim" if st.session_state.contract_cliente_seguradora_checkbox else "Não",
        '_SEGURADORA_': st.session_state.contract_seguradora_input if st.session_state.contract_cliente_seguradora_checkbox else "---",
        '_BOOL_FINANCIADO_': "Sim" if st.session_state.contract_veiculo_financiado_checkbox else "Não",
        '_QTD_PARCELAS_': str(st.session_state.contract_qtd_parcelas_input) if st.session_state.contract_veiculo_financiado_checkbox else "---",
        '_VALOR_PARCELAS_': formatar_valor_financeiro_contrato(st.session_state.contract_valor_parcelas_input) if st.session_state.contract_veiculo_financiado_checkbox and st.session_state.contract_valor_parcelas_input != "---" else "---",
        '_DATA_ULTIMA_PARCELA_': st.session_state.contract_data_ultima_parcela_input.strftime('%d/%m/%Y') if st.session_state.contract_veiculo_financiado_checkbox and st.session_state.contract_data_ultima_parcela_input else "---",
    }

    dados['contract_type'] = st.session_state.contract_tipo_contrato_select
    return dados

def preencher_template_contrato(doc, dados_contrato):
    """Preenche um template DOCX com os dados fornecidos."""
    doc: docx.document.Document
    # Preencher parágrafos (corpo)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, value in dados_contrato['CORPO'].items():
                if key in run.text:
                    run.text = run.text.replace(key, str(value if value is not None else "")).upper()

    # Preencher tabelas
    for table_idx, table in enumerate(doc.tables):
        if dados_contrato['contract_type'] == 'PLANO2' and table_idx == 1: # Contrato PLANO2
            for i in range(len(st.session_state.contract_placas_list)):
                placa_info = st.session_state.contract_placas_list[i]
                row_cells = table.add_row().cells
                row_cells[0].text = f"{placa_info['placa'].upper()} / {placa_info['modelo'].upper()} / {placa_info['marca'].upper()} / ONEBLOCK COM BLOQUEIO."
                row_cells[1].text = "GSM GPRS"
            set_table_borders(table)
        
        elif dados_contrato['contract_type'] != 'PLANO2' and table_idx == 1: # Contrato GSM
            for i in range(len(st.session_state.contract_placas_list)):
                placa_info = st.session_state.contract_placas_list[i]
                row_cells = table.add_row().cells
                row_cells[0].text = placa_info['placa'].upper()
                row_cells[1].text = placa_info['marca'].upper()
                row_cells[2].text = placa_info['modelo'].upper()
                row_cells[3].text = placa_info['rastreador'].upper()
                row_cells[4].text = f"R$ {formatar_valor_financeiro_contrato(placa_info['mensalidade'])}"
            set_table_borders(table)


        
        # Substituição genérica para outros placeholders nas tabelas
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for key, value in dados_contrato['TABELAS'].items():
                            if key in run.text:
                                run.text = run.text.replace(key, str(value if value is not None else "")).upper()
        


def preencher_template_manual(doc, dados_manual):
    """Preenche um template DOCX de manual."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for key, value in dados_manual.items():
                            if key in run.text:
                                run.text = run.text.replace(key, str(value if value is not None else ""))
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, value in dados_manual.items():
                if key in run.text:
                    run.text = run.text.replace(key, str(value if value is not None else ""))