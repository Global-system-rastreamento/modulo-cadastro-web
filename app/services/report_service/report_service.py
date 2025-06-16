import docx
import requests
import json
import datetime
import subprocess
import os
import logging
from time import sleep
import math
import shutil



X_TOKEN_API = 'c0f9e2df-d26b-11ef-9216-0e3d092b76f7'
API_BASE_URL = 'https://api.plataforma.app.br'
POSITIONS_API_URL_TEMPLATE = f'{API_BASE_URL}/report/{{vehicle_id}}/positions'

COMMON_API_HEADERS = {
    'accept': 'application/json, text/plain, */*',
    'accept-language': 'pt-BR,pt;q=0.9,en-US;q=0.8,en;q=0.7',
    'origin': 'https://globalsystem.plataforma.app.br',
    'priority': 'u=1, i',
    'referer': 'https://globalsystem.plataforma.app.br/',
    'sec-ch-ua': '"Not(A:Brand";v="99", "Opera GX";v="118", "Chromium";v="133"',
    'sec-ch-ua-mobile': '?0',
    'sec-ch-ua-platform': '"Windows"',
    'sec-fetch-dest': 'empty',
    'sec-fetch-mode': 'cors',
    'sec-fetch-site': 'same-site',
    'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/133.0.0.0 Safari/537.36 OPR/118.0.0.0',
    'x-token': X_TOKEN_API
}

def get_vehicles(client_id: str) -> list:
    """Busca os veículos de um cliente específico."""
    url = f'{API_BASE_URL}/manager/user/{client_id}/vehicles'
    logging.info(f"Buscando veículos para o cliente ID: {client_id} na URL: {url}")
    try:
        response = requests.get(url, headers=COMMON_API_HEADERS, timeout=30) # Adicionado timeout
        response.raise_for_status()
        vehicles = response.json()
        logging.info(f"Veículos encontrados para o cliente {client_id}: {len(vehicles)}")
        return vehicles
    except requests.exceptions.HTTPError as http_err:
        logging.error(f"Erro HTTP ao buscar veículos para o cliente {client_id}: {http_err}. Resposta: {http_err.response.text if http_err.response else 'N/A'}")
    except requests.exceptions.ConnectionError as conn_err:
        logging.error(f"Erro de conexão ao buscar veículos para o cliente {client_id}: {conn_err}")
    except requests.exceptions.Timeout as timeout_err:
        logging.error(f"Timeout ao buscar veículos para o cliente {client_id}: {timeout_err}")
    except requests.exceptions.RequestException as req_err:
        logging.error(f"Erro geral de requisição ao buscar veículos para o cliente {client_id}: {req_err}")
    except json.JSONDecodeError as json_err:
        logging.error(f"Erro ao decodificar JSON da resposta de veículos para o cliente {client_id}: {json_err}. Resposta: {response.text if 'response' in locals() and response else 'N/A'}")
    return []

def get_tracker_positions(vehicle_id: int, initial_date: str, final_date: str, speed_above: int = None) -> dict:
    """Busca as posições do rastreador para um veículo em um período."""
    url = POSITIONS_API_URL_TEMPLATE.format(vehicle_id=vehicle_id)
    params = {
        'initial_date': initial_date,
        'final_date': final_date,
        'ignition_state': 2,
    }

    if speed_above:
        params['speed_above'] = speed_above

    logging.info(f"Buscando posições para vehicle_id {vehicle_id} de {initial_date} até {final_date} com params: {params}")
    try:
        response = requests.get(url, headers=COMMON_API_HEADERS, params=params, timeout=60)
        response.raise_for_status()
        if response.status_code == 204:
            logging.info(f"Nenhuma posição encontrada (204 No Content) para vehicle_id {vehicle_id} no período.")
            return {"positions": []}

        data = response.json()

        if not isinstance(data, dict) or "positions" not in data:
            logging.warning(f"Resposta inesperada da API de posições para vehicle_id {vehicle_id}: {data}")
            return {"positions": []}

        logging.info(f"Posições encontradas para vehicle_id {vehicle_id}: {len(data.get('positions', []))}")
        return data

    except requests.exceptions.HTTPError as http_err:
        logging.error(f"Erro HTTP ao buscar posições para o veículo {vehicle_id}: {http_err}. URL: {url}. Params: {params}. Resposta: {http_err.response.text if http_err.response else 'N/A'}")
    except requests.exceptions.ConnectionError as conn_err:
        logging.error(f"Erro de conexão ao buscar posições para o veículo {vehicle_id}: {conn_err}. URL: {url}. Params: {params}")
    except requests.exceptions.Timeout as timeout_err:
        logging.error(f"Timeout ao buscar posições para o veículo {vehicle_id}: {timeout_err}. URL: {url}. Params: {params}")
    except requests.exceptions.RequestException as req_err:
        logging.error(f"Erro geral de requisição ao buscar posições para o veículo {vehicle_id}: {req_err}. URL: {url}. Params: {params}")
    except json.JSONDecodeError as json_err:
        logging.error(f"Erro ao decodificar JSON da API de posições para o veículo {vehicle_id}: {json_err}. URL: {url}. Params: {params}. Resposta recebida: {response.text if 'response' in locals() and response else 'N/A'}")
    return {"positions": []}


def get_vehicle_info(vehicle_id):
    url = f"https://api.plataforma.app.br/manager/vehicle/{vehicle_id}?include_mirrored_users=1"
    logging.info(f"Buscando informações para o vehicle_id: {vehicle_id} na URL: {url}")

    headers = {
        "accept": "application/json, text/plain, */*",
        "accept-encoding": "gzip, deflate, br, zstd",
        "accept-language": "pt-BR,pt;q=0.9,en-US;q=0.8,en;q=0.7",
        "origin": "https://globalsystem.plataforma.app.br",
        "referer": "https://globalsystem.plataforma.app.br/",
        "sec-ch-ua": '"Not(A:Brand";v="99", "Opera GX";v="118", "Chromium";v="133"',
        "sec-ch-ua-mobile": "?0",
        "sec-ch-ua-platform": '"Windows"',
        "sec-fetch-dest": "empty",
        "sec-fetch-mode": "cors",
        "sec-fetch-site": "same-site",
        "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/133.0.0.0 Safari/537.36 OPR/118.0.0.0",
        "x-token": "c0f9e2df-d26b-11ef-9216-0e3d092b76f7" # Considerar usar a constante X_TOKEN_API
    }

    try:
        response = requests.get(url, headers=headers, timeout=30) # Adicionado timeout
        response.raise_for_status()
        vehicle_info = response.json()
        logging.info(f"Informações obtidas com sucesso para vehicle_id: {vehicle_id}")
        return vehicle_info
    except requests.exceptions.HTTPError as http_err:
        logging.error(f"Erro HTTP ao buscar informações do veículo {vehicle_id}: {http_err}. Resposta: {http_err.response.text if http_err.response else 'N/A'}")
    except requests.exceptions.ConnectionError as conn_err:
        logging.error(f"Erro de conexão ao buscar informações do veículo {vehicle_id}: {conn_err}")
    except requests.exceptions.Timeout as timeout_err:
        logging.error(f"Timeout ao buscar informações do veículo {vehicle_id}: {timeout_err}")
    except requests.exceptions.RequestException as req_err:
        logging.error(f"Erro geral de requisição ao buscar informações do veículo {vehicle_id}: {req_err}")
    except json.JSONDecodeError as json_err: # Mudado de ValueError para json.JSONDecodeError
        logging.error(f"Resposta não é um JSON válido para vehicle_id {vehicle_id}: {json_err}. Resposta: {response.text if 'response' in locals() and response else 'N/A'}")
        return {"error": "JSONDecodeError", "response_text": response.text if 'response' in locals() and response else 'N/A'} # Retornar um dict em caso de erro
    return {} # Retornar um dict vazio em caso de outros erros

def haversine(lat1, lon1, lat2, lon2):
    """
    Calcula a distância entre dois pontos na Terra (especificados por latitude e longitude)
    usando a fórmula de Haversine.

    Argumentos:
        lat1 (float): Latitude do primeiro ponto em graus.
        lon1 (float): Longitude do primeiro ponto em graus.
        lat2 (float): Latitude do segundo ponto em graus.
        lon2 (float): Longitude do segundo ponto em graus.

    Retorna:
        float: A distância entre os dois pontos em quilômetros.
    """
    # Raio da Terra em quilômetros
    R = 6371.0

    # Converte graus para radianos
    lat1_rad = math.radians(lat1)
    lon1_rad = math.radians(lon1)
    lat2_rad = math.radians(lat2)
    lon2_rad = math.radians(lon2)

    # Diferenças de longitude e latitude
    dlon = lon2_rad - lon1_rad
    dlat = lat2_rad - lat1_rad

    # Fórmula de Haversine
    a = math.sin(dlat / 2)**2 + math.cos(lat1_rad) * math.cos(lat2_rad) * math.sin(dlon / 2)**2
    c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))

    distancia = R * c
    return distancia

def calcular_km_rodados(posicoes):
    """
    Calcula o total de quilômetros rodados a partir de uma lista de posições (latitude, longitude).

    Argumentos:
        posicoes (list): Uma lista de tuplas ou listas, onde cada item contém
                         a latitude e a longitude de um ponto.
                         Exemplo: [(lat1, lon1), (lat2, lon2), (lat3, lon3), ...]

    Retorna:
        float: O total de quilômetros rodados.
               Retorna 0.0 se a lista de posições tiver menos de dois pontos.
    """
    if len(posicoes) < 2:
        return 0.0

    distancia_total = 0.0
    for i in range(len(posicoes) - 1):
        ponto_atual = posicoes[i]
        proximo_ponto = posicoes[i+1]

        lat1, lon1 = ponto_atual
        lat2, lon2 = proximo_ponto

        distancia_segmento = haversine(lat1, lon1, lat2, lon2)
        distancia_total += distancia_segmento

    return int(math.floor(distancia_total))

def parse_datetime_from_api(date_string: str) -> datetime.datetime:
    """Converte string de data da API para objeto datetime."""
    try:
        return datetime.datetime.fromisoformat(date_string.replace("Z", "+00:00").replace("0000", "00:00"))
    except ValueError as e:
        logging.error(f"Erro ao converter string de data da API '{date_string}': {e}")
        return None # Ou levantar a exceção, dependendo do comportamento desejado
    except Exception as e:
        logging.error(f"Erro inesperado ao converter string de data da API '{date_string}': {e}")
        return None


def process_vehicle_positions(
        vehicle_data: dict, 
        processing_date: datetime.datetime, 
        uppertime_threshold: int = 21, 
        lowetime_threshold: int = 5, 
        positions_data: dict = None, 
        ):
    
    """Processa as posições de um veículo para detectar infrações do dia anterior."""
    plate = vehicle_data.get('license_plate', 'PLACA_DESCONHECIDA')
    vehicle_id = vehicle_data.get('id')

    if not vehicle_id:
        logging.warning(f"ID do veículo não encontrado para dados: {vehicle_data}. Pulando.")
        return {}, 'ID do veículo ausente'

    logging.info(f"Iniciando processamento de relatório diário para o veículo: {plate} (ID: {vehicle_id}) para data: {processing_date.strftime('%Y-%m-%d')}")

    previous_day = processing_date - datetime.timedelta(days=1)

    initial_date_dt = datetime.datetime(previous_day.year, previous_day.month, previous_day.day, 0, 0, 0, tzinfo=datetime.timezone.utc)
    final_date_dt = datetime.datetime(previous_day.year, previous_day.month, previous_day.day, 23, 59, 59, tzinfo=datetime.timezone.utc)

    initial_date_str = initial_date_dt.strftime('%Y-%m-%d %H:%M:%S')
    final_date_str = final_date_dt.strftime('%Y-%m-%d %H:%M:%S')

    if not positions_data:
        positions_data = get_tracker_positions(vehicle_id, initial_date_str, final_date_str)

    processed_data = {}
    if not positions_data or "positions" not in positions_data or not positions_data["positions"]:
        logging.warning(f"Sem dados de posições para o veículo {plate} (ID: {vehicle_id}) no período de {initial_date_str} a {final_date_str}.")
        return processed_data, 'Sem dados', positions_data

    positions_ign_on = []
    try:
        positions_ign_on = list(filter(lambda x: x.get('ign') == 1, positions_data['positions']))
    except TypeError as te:
        logging.error(f"Erro ao filtrar posições com ignição ligada para o veículo {plate} (ID: {vehicle_id}). Dados de posições: {positions_data}. Erro: {te}")
        return processed_data, 'Erro ao processar dados de ignição', positions_data


    if not positions_ign_on:
        logging.info(f"Sem dados de posições com ignição ligada para o veículo {plate} (ID: {vehicle_id}).")
        return processed_data, 'Sem dados ign ligado', positions_data

    try:
        trajeto = [(float(pos_trajeto.get('latitude')), float(pos_trajeto.get('longitude'))) for pos_trajeto in positions_ign_on]

        processed_data['totalKm'] = calcular_km_rodados(trajeto)

        processed_data['maxVel'] = max([pos.get('vel', 0) for pos in positions_ign_on if pos.get('vel') is not None] or [0]) # Garante que há pelo menos um 0 para max

        # =====================================================================================================================
        processed_data['hrLigado'] = None
        last_pos_data = None
        first_pos_data = None

        i = 0
        num_positions = len(positions_data.get('positions', []))
        logging.debug(f"Total de posições a serem processadas: {num_positions}")

        while True:
            if i >= num_positions:
                logging.info(f"Loop principal encerrado. Índice i ({i}) atingiu o final da lista de posições ({num_positions}).")
                break

            pos = positions_data['positions'][i]
            logging.debug(f"Loop principal - i: {i}, Posição atual: {pos}")

            if pos.get('ign') == 1:
                logging.info(f"Ignição LIGADA encontrada na posição {i}. Data/Hora: {pos.get('dataHora')}")
                if not first_pos_data:
                    first_pos_data = pos.copy()
                    logging.debug(f"Definido first_pos_data: {first_pos_data}")
                    last_pos_data = first_pos_data.copy()
                    logging.debug(f"Definido last_pos_data inicial como cópia de first_pos_data: {last_pos_data}")


                j = i + 1
                logging.debug(f"Iniciando loop interno. j começa em: {j}")

                while True:
                    process_time = False
                    pos_next = None

                    if j < num_positions:
                        pos_next = positions_data['positions'][j]
                        logging.debug(f"Loop interno - j: {j}, Próxima posição: {pos_next}")

                        if not pos_next.get('ign') == 1:
                            logging.info(f"Ignição DESLIGADA encontrada na próxima posição {j}. Processando tempo.")
                            process_time = True
                        else:
                            logging.debug(f"Ignição continua LIGADA na posição {j}. Atualizando last_pos_data.")
                            last_pos_data = pos_next.copy()
                            process_time = False
                            j += 1
                    else:
                        logging.info(f"Fim da lista de posições ({j}) alcançado no loop interno. Processando tempo com a última ignição LIGADA conhecida.")
                        process_time = True

                    if process_time:
                        logging.debug(f"Condição process_time: {process_time}, first_pos_data: {'Set' if first_pos_data else 'None'}, last_pos_data: {'Set' if last_pos_data else 'None'}")
                        if first_pos_data and last_pos_data:
                            first_dt_str = first_pos_data.get('dataHora')
                            last_dt_str = last_pos_data.get('dataHora')
                            logging.debug(f"Calculando intervalo entre: first_dt_str='{first_dt_str}', last_dt_str='{last_dt_str}'")

                            first_dt: datetime.datetime = parse_datetime_from_api(first_dt_str)
                            last_dt: datetime.datetime = parse_datetime_from_api(last_dt_str)

                            if first_dt and last_dt:
                                if last_dt < first_dt:
                                    logging.warning(f"Timestamp final ({last_dt}) é anterior ao inicial ({first_dt}). Isso pode indicar dados inconsistentes. Intervalo será zero ou negativo.")
                                    interval = datetime.timedelta(0)
                                else:
                                    interval = last_dt - first_dt
                                
                                logging.info(f"Intervalo calculado: {interval}")

                                if processed_data.get('hrLigado') is None:
                                    processed_data['hrLigado'] = interval
                                    logging.debug(f"hrLigado inicializado com: {interval}")
                                else:
                                    processed_data['hrLigado'] += interval
                                    logging.debug(f"hrLigado atualizado para: {processed_data['hrLigado']}")
                            else:
                                logging.error("Não foi possível converter data/hora para calcular o intervalo. Pulando este segmento.")

                            first_pos_data = None
                            last_pos_data = None
                            logging.debug("first_pos_data e last_pos_data resetados para None.")
                            i = j 
                            logging.debug(f"Loop principal 'i' atualizado para {i} (valor de j). Saindo do loop interno.")
                            break 
                        else:
                            logging.warning("Tentativa de processar tempo, mas first_pos_data ou last_pos_data não estão definidos. Verifique a lógica.")
                            i = j
                            logging.debug(f"Loop principal 'i' atualizado para {i} para evitar loop infinito. Saindo do loop interno.")
                            break 
            else:
                logging.debug(f"Ignição DESLIGADA ou ausente na posição {i}. Avançando para a próxima.")
                i += 1

        if processed_data.get('hrLigado') is None:
            logging.info("Nenhum período de ignição ligada foi encontrado ou processado. Definindo hrLigado como 0 segundos.")
            processed_data['hrLigado'] = datetime.timedelta(seconds=0)
        else:
            logging.info(f"Processamento finalizado. Total hrLigado: {processed_data['hrLigado']}")

        processed_data['hrLigado']
        fora_hr_count = 0
        for x in positions_ign_on:
            data_hora_str = x.get('dataHora')
            if not data_hora_str:
                logging.warning(f"Posição sem 'dataHora' para veículo {plate} (ID: {vehicle_id}): {x}")
                continue

            dt_obj = parse_datetime_from_api(data_hora_str)
            if dt_obj:
                if dt_obj.time() > datetime.time(uppertime_threshold, 1, 0) or dt_obj.time() < datetime.time(lowetime_threshold, 0, 0):
                    fora_hr_count += 1
            else:
                logging.warning(f"Não foi possível parsear dataHora '{data_hora_str}' para veículo {plate} (ID: {vehicle_id})")

        processed_data['foraHr'] = 'SIM' if fora_hr_count > 0 else 'NÃO'
        logging.info(f"Processamento para veículo {plate} (ID: {vehicle_id}) concluído. Km: {processed_data['totalKm']}, MaxVel: {processed_data['maxVel']}, ForaHr: {processed_data['foraHr']}")

    except (IndexError, TypeError, KeyError) as e:
        logging.error(f"Erro ao processar dados de posição para veículo {plate} (ID: {vehicle_id}): {e}. Posições com ignição: {len(positions_ign_on)}", exc_info=True)
        return {}, f'Erro nos dados: {e}'

    return processed_data, 'ok', positions_data
    
            
def convert_docx_to_pdf(input_file):
    """
    Converte um arquivo DOCX para PDF usando unoconv.

    Args:
        input_file (str): Caminho para o arquivo DOCX de entrada.
    """
    logging.info(f"Tentando converter {input_file} para PDF.")
    try:
        if not os.path.exists(input_file):
            logging.error(f"Arquivo de entrada para conversão não encontrado: {input_file}")
            raise FileNotFoundError(f"O arquivo {input_file} não foi encontrado.")

        output_file = input_file.replace('.docx', '.pdf')
        if os.path.exists(output_file):
            logging.warning(f"Arquivo de saída {output_file} já existe. Será sobrescrito.")
            try:
                os.remove(output_file)
            except OSError as e:
                logging.error(f"Não foi possível remover o arquivo PDF existente {output_file}: {e}")
                return False, None

        process = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(input_file), input_file],
        check=False, capture_output=True, text=True, timeout=120
    )

        if process.returncode == 0:
            if os.path.exists(output_file) and os.path.getsize(output_file) > 0: # Checa se o PDF foi criado e não está vazio
                logging.info(f"Convertido com sucesso: {input_file} para {output_file}")
                return True, output_file
            else:
                logging.error(f"unoconv retornou 0, mas o arquivo PDF {output_file} não foi criado ou está vazio. STDOUT: {process.stdout}, STDERR: {process.stderr}")
                return False, None
        else:
            logging.error(f"Erro ao converter {input_file} com unoconv. Código de retorno: {process.returncode}")
            logging.error(f"unoconv STDOUT: {process.stdout}")
            logging.error(f"unoconv STDERR: {process.stderr}")
            return False, None

    except subprocess.TimeoutExpired:
        logging.error(f"Timeout ao converter {input_file} para PDF.")
        return False, None
    except subprocess.CalledProcessError as e: # Este erro não será mais levantado com check=False, mas mantido por segurança
        logging.error(f"Erro no subprocesso ao converter {input_file}: {e}. STDOUT: {e.stdout}. STDERR: {e.stderr}")
        return False, None
    except FileNotFoundError: # Já logado acima, mas pode ser pego aqui se unoconv não estiver instalado
        logging.error(f"Comando 'unoconv' não encontrado. Verifique se está instalado e no PATH.")
        return False, None
    except Exception as e:
        logging.error(f"Erro inesperado durante a conversão de {input_file} para PDF: {e}", exc_info=True)
        return False, None


def preencher_doc(template_values, table_lines, date_alvo: datetime.datetime, nome_cliente: str):
    template_path = 'app/services/report_service/Template/Relatório de Quilometragem e Infrações.docx'
    output_dir = "app/services/report_service/Template/preenchido"

    shutil.rmtree(output_dir)
    os.makedirs(output_dir, exist_ok=True) # Garante que o diretório de saída exista

    logging.info(f"Preenchendo documento para cliente '{nome_cliente}' para data {date_alvo.strftime('%d-%m-%Y')}.")
    try:
        if not os.path.exists(template_path):
            logging.error(f"Template DOCX não encontrado em: {template_path}")
            return False, None

        doc = docx.Document(template_path)

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                for k, v in template_values.items():
                    if str(k) in run.text:
                        run.text = run.text.replace(str(k), str(v)) # Convertido k para str para segurança

        if not doc.tables:
            logging.error(f"Nenhuma tabela encontrada no template: {template_path}")
            return False, None

        table = doc.tables[0]

        # Adiciona linhas apenas se table_lines não estiver vazio
        if table_lines:
            for _ in range(len(table_lines)):
                try:
                    table.add_row()
                except Exception as e: # Captura erro específico ao adicionar linha, se houver
                    logging.error(f"Erro ao adicionar linha na tabela do DOCX: {e}", exc_info=True)
                    return False, None

            # Começa a preencher a partir da segunda linha (índice 1), assumindo que a primeira é cabeçalho
            # E garante que o número de linhas adicionadas corresponda a table_lines
            if len(table.rows) < 1 + len(table_lines):
                logging.error(f"Número de linhas na tabela ({len(table.rows)}) é menor que o esperado após adicionar {len(table_lines)} linhas.")
                return False, None

            for i, row_new_line_data in enumerate(table_lines):
                row_table = table.rows[i + 1] # Acessa a linha correta (pós-cabeçalho)
                if len(row_table.cells) != len(row_new_line_data):
                    logging.warning(f"Número de células na linha {i+1} da tabela ({len(row_table.cells)}) não corresponde aos dados fornecidos ({len(row_new_line_data)}). Pulando preenchimento desta linha.")
                    continue
                for cell, value in zip(row_table.cells, row_new_line_data):
                    # Limpa parágrafos existentes na célula antes de adicionar novo, se necessário
                    # for p in cell.paragraphs:
                    # p.clear() # Ou uma forma de remover o parágrafo se estiver vazio
                    cell.text = str(value) # Mais simples para preencher células

        # Sanitizar nome do cliente para nome de arquivo
        safe_client_name = "".join(c if c.isalnum() or c in " _-" else "_" for c in nome_cliente)
        path = os.path.join(output_dir, f"Relatório {date_alvo.strftime('%d-%m-%Y')} de Quilometragem e Infrações {safe_client_name}.docx")

        doc.save(path)
        logging.info(f"Documento DOCX salvo em: {path}")
        return True, path
    except FileNotFoundError: # Caso o template_path não seja encontrado (já verificado, mas bom ter)
        logging.error(f"Arquivo template não encontrado em '{template_path}'.", exc_info=True)
        return False, None
    except Exception as e:
        logging.error(f"Erro ao preencher ou salvar o documento DOCX para '{nome_cliente}': {e}", exc_info=True)
        return False, None

def process_report(client_config):
    logging.info("Iniciando execução do processo de relatório diário.")
    date_alvo = datetime.datetime.now()

    client_name = client_config.get('name')
    client_id = client_config.get('cod')
    vel_threshold = client_config.get('vel_threshold', 100)
    uppertime_threshold = client_config.get('uppertime_threshold', 21)
    lowertime_threshold = client_config.get('lowertime_threshold', 5)
    plates_to_not_do = client_config.get('plates_to_not_do', [])

    if not client_id:
        logging.warning(f"Cliente '{client_name}' não possui 'cod' configurado. Pulando.")
        return

    logging.info(f"Processando cliente: {client_name} (ID: {client_id})")

    try:
        client_vehicles = get_vehicles(client_id)
        if not client_vehicles:
            logging.warning(f"Nenhum veículo encontrado para o cliente {client_name} (ID: {client_id}). Pulando este cliente.")
            return

        nome_cliente_proprietario = client_vehicles[0].get('owner', {}).get('name', client_name) # Usa nome do cliente se proprietário não encontrado

        template_values = {
            '_NOME_CLIENTE_': nome_cliente_proprietario,
            '_PERIODO_ANALISADO_': (date_alvo - datetime.timedelta(days=1)).strftime('%d/%m/%Y'),
            '_TOTAL_VEICULOS_': len(client_vehicles),
            '_TOTAL_RODADO_': 0,
            '_EXCESSO_VEL_': 0,
            '_VEICULO_FORA_HOR_': 0,
            '_VEICULO_SEM_MOV_': 0,
            '_VEICULO_SEM_HIST_': 0
        }
        table_lines = []

        all_positions_data = []

        for vehicle in client_vehicles:
            if plates_to_not_do and vehicle.get('license_plate') in plates_to_not_do:
                logging.info(f"Veículo {vehicle.get('license_plate', 'N/A')} (ID: {vehicle.get('id')}) está na lista de exclusão para o cliente {client_name}. Pulando.")
                return

            logging.info(f"Processando veículo: {vehicle.get('license_plate', 'N/A')} (ID: {vehicle.get('id')}) para cliente {client_name}")
            processed_data, info, positions_data = process_vehicle_positions(vehicle, date_alvo, uppertime_threshold, lowertime_threshold) # Passando vel_threshold

            all_positions_data.append(positions_data)
            
            placa = vehicle.get('license_plate', 'N/A')
            marca = vehicle.get('brand', 'N/A')
            modelo = vehicle.get('model', 'N/A')

            if info == 'Sem dados':
                template_values['_VEICULO_SEM_HIST_'] += 1
                table_lines.append([placa, marca, modelo, 'Sem Dados', 'Sem Dados', 'Sem Dados', 'Sem Dados'])
            elif info == 'Sem dados ign ligado':
                template_values['_VEICULO_SEM_MOV_'] += 1
                table_lines.append([placa, marca, modelo, 0, '0 H', 0, 'NÃO']) 
            elif isinstance(processed_data, dict) and processed_data and info == 'ok':
                if processed_data.get('maxVel', 0) > vel_threshold:
                    template_values['_EXCESSO_VEL_'] += 1

                if processed_data.get('foraHr') == 'SIM':
                    template_values['_VEICULO_FORA_HOR_'] += 1

                template_values['_TOTAL_RODADO_'] += processed_data.get('totalKm', 0)
                table_lines.append([placa, marca, modelo, processed_data.get('totalKm', 0), str(processed_data.get('hrLigado', 0))[:-3] + ' H', processed_data.get('maxVel', 0), processed_data.get('foraHr', 'NÃO')])
            else:
                logging.warning(f"Processamento do veículo {placa} retornou estado inesperado: {info}, dados: {processed_data}")
                template_values['_VEICULO_SEM_HIST_'] += 1
                table_lines.append([placa, marca, modelo, 'Erro Proc.', 'Erro Proc.', 'Erro Proc.', 'Erro Proc.'])


        logging.info(f"Dados agregados para {client_name}: {template_values}")
        success_doc, path_docx = preencher_doc(template_values, table_lines, date_alvo, nome_cliente_proprietario)

        if success_doc and path_docx:
            logging.info(f"Documento DOCX criado: {path_docx}")
            success_pdf, path_pdf = convert_docx_to_pdf(path_docx)

            if success_pdf and path_pdf and os.path.exists(path_pdf):
                logging.info(f"Documento PDF criado: {path_pdf}")

                return path_pdf
            else:
                logging.error(f"Falha ao converter DOCX para PDF ou PDF não encontrado para {client_name}. DOCX: {path_docx}")
        else:
            logging.error(f"Falha ao criar o documento DOCX para {client_name}.")

    except Exception as e:
        logging.error(f"Erro inesperado ao processar o cliente {client_name} (ID: {client_id}): {e}", exc_info=True)
        # Notificar admin sobre a falha no processamento deste cliente

    logging.info("Processo de relatório diário concluído.")